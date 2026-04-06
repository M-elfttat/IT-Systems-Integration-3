import re
from pathlib import Path

import docx


def main() -> None:
    ws = Path(r"c:\Users\ahmed\OneDrive\Desktop\System integration project")
    md_path = ws / "phase1_submission" / "Phase_1_Report_Updated.md"
    out_path = ws / "phase1_submission" / "Phase_1_Report_Updated.docx"

    md = md_path.read_text(encoding="utf-8")
    lines = md.splitlines()

    d = docx.Document()

    for raw in lines:
        line = raw.rstrip("\n").rstrip("\r")
        if not line.strip():
            continue

        # headings: ##..###### -> level 1..4
        m = re.match(r"^(#{2,6})\s+(.*)$", line)
        if m:
            hashes, text = m.group(1), m.group(2)
            level = max(1, min(4, len(hashes) - 1))
            d.add_heading(text.strip(), level=level)
            continue

        if line.strip() == "---":
            d.add_paragraph(" ")
            continue

        # Bold label lines like: **Project Title:** ...
        if line.strip().startswith("**") and "**" in line.strip()[2:]:
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            p = d.add_paragraph()
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part.strip("*"))
                    r.bold = True
                else:
                    p.add_run(part)
            continue

        if line.lstrip().startswith("- "):
            d.add_paragraph(line.lstrip()[2:], style="List Bullet")
            continue

        d.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(out_path))
    print(out_path)


if __name__ == "__main__":
    main()

